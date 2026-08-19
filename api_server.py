"""
FastAPI server for DataFactZ HR Policy Assistant RAG Chatbot
Connects the React web app to the RAG backend
"""

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from typing import List, Optional
from datetime import datetime
import json
import os
from pathlib import Path

# Import RAG modules
from improved_chunking import ImprovedChunker
from sklearn.metrics.pairwise import cosine_similarity
import numpy as np

# Lazy load sentence_transformers (not available in production)
try:
    from sentence_transformers import SentenceTransformer
    EMBEDDINGS_AVAILABLE = True
except ImportError:
    EMBEDDINGS_AVAILABLE = False
    SentenceTransformer = None

app = FastAPI(
    title="DataFactZ RAG API",
    description="Backend API for HR Policy Assistant",
    version="1.0.0"
)

# Enable CORS
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Initialize models
embedding_model = None
try:
    if EMBEDDINGS_AVAILABLE and SentenceTransformer:
        embedding_model = SentenceTransformer('all-MiniLM-L6-v2')
        print("✓ Embedding model loaded")
    else:
        print("⚠ Embedding model not available (using keyword search fallback)")
except Exception as e:
    print(f"⚠ Could not load embedding model: {e}")

chunker = ImprovedChunker(chunk_size=300, overlap=50)
print("✓ Chunker initialized")

# In-memory storage (replace with database for production)
documents_store = []
chunks_store = []
embeddings_store = []
full_documents = {}  # Store full document text for viewing

# Pydantic models
class Citation(BaseModel):
    document_title: str
    document_id: str  # Actual document ID for API calls
    section_path: str
    excerpt: str
    relevance_score: float
    chunk_index: Optional[int] = None  # Index for retrieving full context

class QueryRequest(BaseModel):
    query: str
    conversation_history: Optional[List[dict]] = None

class QueryResponse(BaseModel):
    answer: str
    citations: List[Citation]
    confidence: str
    generation_time_seconds: float

class DocumentInfo(BaseModel):
    id: str
    filename: str
    title: str
    description: str
    chunk_count: int
    token_count: int
    indexed_at: str
    status: str

class DocumentsResponse(BaseModel):
    documents: List[DocumentInfo]

# Routes

@app.get("/health")
async def health_check():
    """Check API health"""
    return {
        "status": "healthy",
        "timestamp": datetime.now().isoformat(),
        "models_loaded": True,
    }

async def ensure_documents_indexed():
    """Ensure documents are indexed, auto-reindex if needed (for multi-worker environments)"""
    global chunks_store, embeddings_store, documents_store, full_documents

    try:
        data_folder = Path("sample_data")
        if data_folder.exists():
            # Get ALL supported file types (txt, pdf, docx)
            policy_files = list(data_folder.glob("*.txt")) + list(data_folder.glob("*.pdf")) + list(data_folder.glob("*.docx"))
            if policy_files:
                chunks_store = []
                embeddings_store = []
                documents_store = []
                full_documents = {}

                for file_path in sorted(policy_files):
                    print(f"[STARTUP] Indexing: {file_path.name}")
                    full_text = ""
                    file_ext = file_path.suffix.lower()

                    try:
                        # Extract text based on file type
                        if file_ext == '.txt':
                            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                                full_text = f.read()
                        elif file_ext == '.pdf':
                            try:
                                import PyPDF2
                                with open(file_path, 'rb') as pdf_file:
                                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                                    for page_num in range(len(pdf_reader.pages)):
                                        page = pdf_reader.pages[page_num]
                                        full_text += page.extract_text() + "\n"
                            except Exception as pdf_err:
                                print(f"PDF extraction error for {file_path.name}: {pdf_err}")
                                continue
                        elif file_ext == '.docx':
                            try:
                                from docx import Document
                                doc = Document(file_path)
                                for paragraph in doc.paragraphs:
                                    full_text += paragraph.text + "\n"
                            except Exception as docx_err:
                                print(f"DOCX extraction error for {file_path.name}: {docx_err}")
                                continue

                        if not full_text.strip():
                            print(f"[STARTUP] Skipping empty file: {file_path.name}")
                            continue

                        # Use unique doc_id that includes file extension to avoid collisions
                        doc_id = file_path.stem + "_" + file_ext.strip(".")
                        full_documents[doc_id] = full_text

                        # Use detect_sections for consistent chunking
                        chunks = chunker.detect_sections(full_text)
                        chunk_dicts = []
                        for idx, chunk_text in enumerate(chunks, 1):
                            token_count = chunker.count_tokens(chunk_text)
                            chunk_dicts.append({
                                'text': chunk_text,
                                'chunk_number': idx,
                                'document_title': doc_id,
                                'section_path': f'Section {idx}',
                                'token_count': token_count,
                            })

                        for chunk_dict in chunk_dicts:
                            chunks_store.append(chunk_dict)

                        chunk_texts = [c['text'] for c in chunk_dicts]
                        if embedding_model:
                            chunk_embeddings = embedding_model.encode(chunk_texts, show_progress_bar=False)
                        else:
                            chunk_embeddings = np.zeros((len(chunk_texts), 384))
                        embeddings_store.extend(chunk_embeddings)

                        total_tokens = sum(c['token_count'] for c in chunk_dicts)
                        documents_store.append(DocumentInfo(
                            id=doc_id,
                            filename=file_path.name,
                            title=file_path.stem.replace('_', ' ').title(),
                            description=f"HR policy document with {len(chunk_dicts)} chunks",
                            chunk_count=len(chunk_dicts),
                            token_count=total_tokens,
                            indexed_at=datetime.now().isoformat(),
                            status="indexed"
                        ))
                        print(f"[STARTUP] ✓ Indexed {file_path.name}: {len(chunk_dicts)} chunks, {total_tokens} tokens")
                    except Exception as file_err:
                        print(f"[STARTUP] Error processing {file_path.name}: {file_err}")
                        continue

    except Exception as e:
        print(f"Auto-reindex error: {e}")

@app.post("/api/query", response_model=QueryResponse)
async def query_documents(request: QueryRequest):
    """
    Process a natural language query against indexed documents

    Returns grounded answer with citations
    """
    if not request.query.strip():
        raise HTTPException(status_code=400, detail="Query cannot be empty")

    if not chunks_store or not embeddings_store:
        raise HTTPException(
            status_code=503,
            detail="No documents indexed. Sample data folder may be missing."
        )

    try:
        import time
        start_time = time.time()
        print(f"[QUERY] Processing: {request.query[:50]}...")

        # Find similar chunks using embeddings or keyword matching
        print(f"[QUERY] Using embedding model: {embedding_model is not None}")
        if embedding_model and embeddings_store:
            # Embed query using sentence transformers
            query_embedding = embedding_model.encode(request.query, show_progress_bar=False)
            similarities = cosine_similarity([query_embedding], embeddings_store)[0]
            top_indices = np.argsort(similarities)[::-1][:5]
        else:
            # Fallback: improved keyword matching with word stem matching
            query_words = request.query.lower().split()
            similarities = []

            # Filter out common stop words
            stop_words = {'the', 'a', 'an', 'and', 'or', 'is', 'to', 'of', 'in', 'at', 'for', 'by', 'this', 'that', 'are', 'was', 'be', 'been', 'have', 'has', 'do', 'does', 'did', 'will', 'would', 'could', 'should', 'may', 'might', 'must', 'can', 'as', 'with', 'from', 'on', 'it', 'its'}
            important_words = [w for w in query_words if w not in stop_words]

            for chunk in chunks_store:
                chunk_text = chunk['text'].lower()
                chunk_words = chunk_text.split()

                # Score: count how many query words appear in chunk
                score = 0
                for qw in important_words:
                    # Exact word match (highest priority)
                    if qw in chunk_words:
                        score += 40
                    # Substring match or word stem - check if any word starts with query word
                    elif any(word.startswith(qw[:3]) if len(qw) >= 3 else False for word in chunk_words):
                        score += 20
                    # Any appearance in text (low priority)
                    elif qw in chunk_text:
                        score += 10

                # Normalize score to 0-1 range
                max_score = len(important_words) * 40 if important_words else 1
                similarity = min(score / max_score, 1.0) if max_score > 0 else 0
                similarities.append(similarity)

            similarities = np.array(similarities)
            top_indices = np.argsort(similarities)[::-1][:5]

        print(f"[QUERY] Found top {len(top_indices)} similar chunks")
        # Extract citations (very low threshold for newly uploaded documents)
        citations = []
        for idx in top_indices:
            if similarities[idx] > 0.15:  # Very low threshold to include all relevant results
                chunk = chunks_store[idx]
                # Get more context from the chunk
                excerpt = chunk['text'][:300] if len(chunk['text']) > 300 else chunk['text']

                # Document ID is already stored in chunk's document_title
                doc_id = chunk.get('document_title', 'Unknown')

                # Get document info for display title
                doc_info = next((d for d in documents_store if d.id == doc_id), None)
                display_title = doc_info.title if doc_info else doc_id.replace('_', ' ').title()

                citations.append(Citation(
                    document_title=display_title,
                    document_id=doc_id,
                    section_path=chunk.get('section_path', ''),
                    excerpt=excerpt,
                    relevance_score=float(similarities[idx]),
                    chunk_index=idx
                ))

        # Determine confidence level (adjusted for lower scores)
        if len(citations) > 0:
            top_score = citations[0].relevance_score
            if top_score > 0.5:
                confidence = "HIGH"
            elif top_score > 0.3:
                confidence = "MEDIUM"
            else:
                confidence = "LOW"
        else:
            confidence = "LOW"

        # Build answer with bullet point formatting
        if citations:
            # Format each citation as a section
            formatted_sections = []
            for i, c in enumerate(citations[:3], 1):
                # Extract key points from excerpt
                excerpt_lines = c.excerpt.strip().split('\n')
                points = [line.strip() for line in excerpt_lines if line.strip() and len(line.strip()) > 10]

                section = f"## {i}️⃣ **{c.document_title}**\n\n**Section:** {c.section_path}\n\n**Key Points:**\n"
                for point in points[:5]:  # Limit to 5 points per section
                    section += f"• {point}\n"
                section += f"\n**Match Score:** {int(c.relevance_score * 100)}%"
                formatted_sections.append(section)

            answer = "## 📋 **ANSWER**\n\n" + "\n\n---\n\n".join(formatted_sections)
        else:
            answer = "❌ I couldn't find relevant information in the policy documents.\n\n**Try:** Rephrase your question or ask about specific topics like vacation, remote work, or conduct."

        generation_time = time.time() - start_time
        print(f"[QUERY] Generated answer in {generation_time:.2f}s with {len(citations)} citations")

        return QueryResponse(
            answer=answer,
            citations=citations,
            confidence=confidence,
            generation_time_seconds=generation_time
        )

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Query processing error: {str(e)}")

@app.get("/api/documents", response_model=DocumentsResponse)
async def get_documents():
    """Get list of indexed documents"""
    return DocumentsResponse(documents=documents_store)

@app.post("/api/reindex")
async def reindex_documents():
    """Reindex all documents in sample_data folder - supports all file types"""
    global chunks_store, embeddings_store, documents_store, full_documents

    try:
        await ensure_documents_indexed()

        return {
            "status": "success",
            "documents_indexed": len(documents_store),
            "total_chunks": len(chunks_store),
            "total_tokens": sum(d.token_count for d in documents_store)
        }

    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Reindex error: {str(e)}")

@app.delete("/api/documents/{doc_id}")
async def delete_document(doc_id: str):
    """Delete a document"""
    global documents_store

    documents_store = [d for d in documents_store if d.id != doc_id]
    return {"status": "deleted", "document_id": doc_id}

@app.post("/api/documents/upload")
async def upload_document(file: UploadFile = File(...)):
    """Upload and index a new document"""
    global chunks_store, embeddings_store, documents_store, full_documents

    try:
        # Save file
        save_path = Path("sample_data") / file.filename
        save_path.parent.mkdir(parents=True, exist_ok=True)

        with open(save_path, "wb") as f:
            content = await file.read()
            f.write(content)

        # Extract text based on file type
        file_ext = save_path.suffix.lower()
        full_text = ""

        print(f"Uploading file: {file.filename} ({file_ext})")

        if file_ext == '.txt':
            print("Reading .txt file...")
            with open(save_path, 'r', encoding='utf-8', errors='ignore') as f:
                full_text = f.read()

        elif file_ext == '.pdf':
            print("Extracting text from PDF...")
            try:
                import PyPDF2
                with open(save_path, 'rb') as pdf_file:
                    pdf_reader = PyPDF2.PdfReader(pdf_file)
                    for page_num in range(len(pdf_reader.pages)):
                        page = pdf_reader.pages[page_num]
                        full_text += page.extract_text() + "\n"
            except Exception as pdf_err:
                print(f"PDF extraction error: {pdf_err}")
                raise Exception(f"Failed to extract PDF text: {str(pdf_err)}")

        elif file_ext == '.docx':
            print("Extracting text from DOCX...")
            try:
                from docx import Document
                doc = Document(save_path)
                for paragraph in doc.paragraphs:
                    full_text += paragraph.text + "\n"
            except Exception as docx_err:
                print(f"DOCX extraction error: {docx_err}")
                raise Exception(f"Failed to extract DOCX text: {str(docx_err)}")
        else:
            raise Exception(f"Unsupported file type: {file_ext}. Supported: .txt, .pdf, .docx")

        if not full_text.strip():
            raise Exception("File is empty or could not be read")

        # Use unique doc_id that includes file extension to avoid collisions
        doc_id = save_path.stem + "_" + file_ext.strip(".")
        full_documents[doc_id] = full_text

        # Chunk the text directly without using process_document
        print("Chunking document...")
        chunks = chunker.detect_sections(full_text)

        # Format chunks like process_document does
        chunk_dicts = []
        for idx, chunk_text in enumerate(chunks, 1):
            token_count = chunker.count_tokens(chunk_text)
            chunk_dicts.append({
                'text': chunk_text,
                'chunk_number': idx,
                'document_title': doc_id,
                'section_path': f'Section {idx}',
                'token_count': token_count,
            })

        chunks = chunk_dicts

        # Store chunks
        for chunk_dict in chunks:
            chunks_store.append(chunk_dict)

        # Create embeddings
        chunk_texts = [c['text'] for c in chunks]
        if embedding_model:
            chunk_embeddings = embedding_model.encode(chunk_texts, show_progress_bar=False)
        else:
            chunk_embeddings = np.zeros((len(chunk_texts), 384))
        embeddings_store.extend(chunk_embeddings)

        # Add document info
        total_tokens = sum(c['token_count'] for c in chunks)
        doc_info = DocumentInfo(
            id=doc_id,
            filename=file.filename,
            title=save_path.stem.replace('_', ' ').title(),
            description="Newly uploaded document",
            chunk_count=len(chunks),
            token_count=total_tokens,
            indexed_at=datetime.now().isoformat(),
            status="indexed"
        )
        documents_store.append(doc_info)

        return {
            "status": "uploaded",
            "filename": file.filename,
            "chunks": len(chunks),
            "tokens": total_tokens
        }

    except Exception as e:
        print(f"Upload error details: {type(e).__name__}: {str(e)}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=f"Upload error: {str(e)}")

@app.get("/api/documents/{doc_id}/content")
async def get_document_content(doc_id: str, preview_only: bool = False):
    """Retrieve document text for viewing (optimized for large documents)"""
    print(f"[DOCUMENT] Requested doc_id: '{doc_id}'")
    print(f"[DOCUMENT] Available doc_ids in full_documents: {list(full_documents.keys())}")

    if doc_id not in full_documents:
        print(f"[DOCUMENT] ERROR: doc_id '{doc_id}' not found in full_documents")
        raise HTTPException(status_code=404, detail=f"Document '{doc_id}' not found")

    content = full_documents[doc_id]

    # For very large documents, optionally send preview
    if preview_only and len(content) > 50000:
        content = content[:50000] + "\n\n[Preview truncated - document is too large to display fully]"

    return {
        "document_id": doc_id,
        "content": content,
        "title": doc_id.replace('_', ' ').title(),
        "is_preview": preview_only and len(content) > 50000
    }

@app.get("/api/chunks/{chunk_index}")
async def get_chunk_context(chunk_index: int):
    """Retrieve a specific chunk with surrounding context"""
    if chunk_index < 0 or chunk_index >= len(chunks_store):
        raise HTTPException(status_code=404, detail="Chunk not found")

    chunk = chunks_store[chunk_index]

    # Get surrounding chunks for context
    context_before = chunks_store[chunk_index - 1]['text'] if chunk_index > 0 else ""
    context_after = chunks_store[chunk_index + 1]['text'] if chunk_index < len(chunks_store) - 1 else ""

    return {
        "chunk_index": chunk_index,
        "document_title": chunk.get('document_title', 'Unknown'),
        "section_path": chunk.get('section_path', ''),
        "text": chunk['text'],
        "context_before": context_before,
        "context_after": context_after,
        "chunk_number": chunk.get('chunk_number', 0),
        "token_count": chunk.get('token_count', 0)
    }

@app.get("/docs", include_in_schema=True)
async def get_docs():
    """API documentation"""
    return {
        "title": "DataFactZ RAG API",
        "version": "1.0.0",
        "endpoints": [
            "GET /health - Check API health",
            "POST /api/query - Query documents",
            "GET /api/documents - List indexed documents",
            "GET /api/documents/{doc_id}/content - Get full document",
            "GET /api/chunks/{chunk_index} - Get specific chunk",
            "POST /api/reindex - Reindex all documents",
            "DELETE /api/documents/{doc_id} - Delete document",
            "POST /api/documents/upload - Upload new document",
        ]
    }

@app.on_event("startup")
async def startup_event():
    """Auto-reindex documents on startup - calls ensure_documents_indexed for all file types"""
    print("Starting DataFactZ RAG API Server...")
    print("Available at http://localhost:8000")
    print("API docs at http://localhost:8000/docs")
    print("Indexing sample documents on startup...")
    await ensure_documents_indexed()
    print(f"✓ Indexed {len(documents_store)} documents with {len(chunks_store)} chunks")

if __name__ == "__main__":
    import uvicorn
    print("Starting DataFactZ RAG API Server...")
    print("Available at http://localhost:8000")
    print("API docs at http://localhost:8000/docs")

    uvicorn.run(app, host="0.0.0.0", port=8000)
